"""SharePoint/OneDrive live tools, SCOPED to the folder the source agent named.
See connectors/confluence.py's module docstring for the shared build_tools
contract every connector module in this package follows.

Purpose-built rather than the generic REST fallback for two reasons:
  1. Scope. An app credential with Sites.Read.All can reach every site in the
     tenant (99 in the test tenant). The source Copilot agent pointed at ONE
     folder, so the migrated agent must be confined to that folder — a tool
     that cannot express a wider path is a stronger guarantee than an
     instruction asking it not to wander.
  2. Reading files. Graph returns raw bytes; the model needs text. Extraction
     (pdf/docx/xlsx) has to happen in the container.

`sharepoint_list_lists` covers the connector's `GetAllTables` operation ("Get all
lists and libraries"). It is here rather than in VENDOR_BINDINGS because the source
operation takes a `dataset` (a site URL) and Graph addresses sites by id or by the
`{host}:{server-relative-path}:` form — a transform no URL template can express.
`_resolve_scope` already does exactly that conversion for the file tools.

Chosen by measured demand, not by swagger order: of 340 operations across the three
proxy-only Microsoft connectors, `GetAllTables` is the ONE that any staged agent
actually calls (`_diag_ms_op_usage.ts`, 131 agents, 2026-08-19). Mapping in swagger
order would have spent weeks before reaching it.
"""


def build_tools(conn, secret, mint_token, auth_header, fill):
    scope_uri = conn.get("scopeUri") or ""

    def _graph(path: str, token: str, raw: bool = False):
        import json as _json
        import urllib.request
        req = urllib.request.Request(
            f"https://graph.microsoft.com/v1.0{path}",
            headers={"Authorization": f"Bearer {token}", "Accept": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        return data if raw else _json.loads(data.decode("utf-8"))

    def _resolve_scope(token: str):
        """Turn the SharePoint URL into (siteId, folderPath). Cached per container."""
        import urllib.parse
        if not scope_uri:
            return None, ""
        p = urllib.parse.urlparse(scope_uri)
        host = p.netloc
        parts = [urllib.parse.unquote(x) for x in p.path.split("/") if x]
        # Three real shapes, not one:
        #   /sites/<name>/<library>/<folders...>        team site
        #   /personal/<user>/<library>/<folders...>     OneDrive-for-Business ("<tenant>-my" host)
        #   /<library>/<folders...>                     root site
        #
        # "personal" was missing, and the failure was quiet in the worst way: parts[0] is
        # "personal", so it fell to the root-site branch, resolved the -my host's ROOT site, and
        # then looked for "personal/<user>/..." as a FOLDER inside that site's default drive.
        # Graph answers 404 and the agent reports "the default SharePoint folder is not
        # configured", which sounds like a setup mistake rather than a parsing bug. Confirmed
        # live 2026-08-22 against
        #   https://filefuze-my.sharepoint.com/personal/alex_filefuze_co/Documents/...
        # Teams chat attachments always live on a personal site, so any agent grounded on a
        # file someone shared in a Teams chat hit this.
        if parts and parts[0].lower() in ("sites", "personal") and len(parts) >= 2:
            site_path = f"/{parts[0].lower()}/{parts[1]}"
            rest = parts[2:]
        else:
            site_path = ""
            rest = parts
        site = _graph(f"/sites/{host}:{site_path}" if site_path else f"/sites/{host}", token)
        # Drop the document-library segment ("Shared Documents"); what remains is the
        # folder path inside the default drive.
        if rest and rest[0].lower() in ("shared documents", "documents", "shared%20documents"):
            rest = rest[1:]
        return site.get("id"), "/".join(rest)

    def _scoped_path(folder: str, user_path: str) -> str:
        """Join a model-supplied relative path onto the scoped folder, rejecting
        any attempt to escape it. The tool signature only documents "a path inside
        the connected folder" as a convention — nothing stops the calling model
        from passing "../../other-site" instead, so this is enforced here rather
        than trusted from the docstring.

        Checked by containment (resolved candidate must equal `folder` or sit
        under it), not just by pattern-matching for ".." — a bare ".." against a
        single-segment folder normalizes straight to ".", which slips past a
        ".."-prefix check while still resolving outside the intended scope.
        """
        import posixpath
        folder = (folder or "").strip("/")
        candidate = posixpath.normpath(posixpath.join(folder, user_path.strip("/").lstrip("/")))
        candidate = "" if candidate == "." else candidate.strip("/")
        escapes = (
            posixpath.isabs(user_path)
            or candidate == ".."
            or candidate.startswith("../")
            or (folder and candidate != folder and not candidate.startswith(folder + "/"))
        )
        if escapes:
            raise ValueError(f"path '{user_path}' escapes the connected SharePoint folder")
        return candidate

    def _enc(p: str) -> str:
        """Percent-encode a drive path for a Graph URL, keeping "/" as the separator.

        SharePoint paths routinely contain spaces and brackets — "Microsoft Teams Chat Files",
        "Ben file 2[1]_1779290909_6257.pdf" — and urllib refuses the request outright rather
        than encoding for you:
            URL can't contain control characters ... (found at least ' ')
        Confirmed live 2026-08-22 on a deployed agent: sharepoint_list_files failed on the very
        first call against a real personal site, so EVERY SharePoint tool was unusable against
        any tenant whose folder names have spaces, which is to say all of them.
        """
        from urllib.parse import quote
        return quote(p, safe="/")

    def sharepoint_list_files(subfolder: str = "") -> dict:
        """List files and folders in the company's SharePoint folder this agent is
        connected to. Only this folder and things inside it are accessible.

        Args:
            subfolder: optional path INSIDE the connected folder, e.g. "Reports/2026".

        Returns:
            dict with `items` (name, isFolder, size, lastModified, id) or `error`.
        """
        try:
            token = mint_token(fill)
            site_id, folder = _resolve_scope(token)
            if not site_id:
                return {"error": "no SharePoint scope configured for this agent"}
            path = _scoped_path(folder, subfolder)
            # A scope can legitimately BE a single file: the orchestrator connects whatever the
            # Copilot author named, and some agents name one document. Asking for the children
            # of a file is meaningless, so ask what the item IS first and answer accordingly —
            # returning the file itself, which is what the caller can actually act on.
            if path:
                item = _graph(f"/sites/{site_id}/drive/root:/{_enc(path)}", token)
                if "file" in item:
                    return {
                        "folder": scope_uri,
                        "subfolder": subfolder,
                        "items": [{
                            "name": item.get("name"),
                            "isFolder": False,
                            "size": item.get("size"),
                            "lastModified": item.get("lastModifiedDateTime"),
                            "id": item.get("id"),
                        }],
                        "count": 1,
                        "note": "This agent is connected to a single file, not a folder. "
                                "Use sharepoint_read_file to read it.",
                    }
            url = (
                f"/sites/{site_id}/drive/root:/{_enc(path)}:/children"
                if path else f"/sites/{site_id}/drive/root/children"
            )
            data = _graph(url, token)
        except Exception as e:  # noqa: BLE001
            return {"error": f"SharePoint list failed: {e}"}
        items = [{
            "name": i.get("name"),
            "isFolder": "folder" in i,
            "size": i.get("size"),
            "lastModified": i.get("lastModifiedDateTime"),
            "id": i.get("id"),
        } for i in data.get("value", [])]
        return {"folder": scope_uri, "subfolder": subfolder, "items": items, "count": len(items)}

    def sharepoint_read_file(file_path: str) -> dict:
        """Read the TEXT CONTENT of a file in the connected SharePoint folder, so you
        can answer questions about what a document says.

        Supports .txt .md .csv .json .log .xml, PDF, Word (.docx) and Excel (.xlsx).
        Images and other binary formats cannot be read.

        Args:
            file_path: file name, or path inside the connected folder,
                e.g. "daily_queries.txt" or "Reports/Q1.pdf".

        Returns:
            dict with `text` (extracted, possibly truncated) or `error`.
        """
        import io
        MAX_BYTES = 20 * 1024 * 1024
        MAX_CHARS = 60000
        try:
            token = mint_token(fill)
            site_id, folder = _resolve_scope(token)
            if not site_id:
                return {"error": "no SharePoint scope configured for this agent"}
            # When the connected scope IS a single file, the caller's file_path must not be
            # joined onto it — that builds ".../Ben file.pdf/Ben file.pdf" and 404s. The model
            # has no way to know the difference: sharepoint_list_files correctly reported the
            # one file, so it asks to read it by name, which is the obvious next move.
            # Confirmed live 2026-08-22, the failure immediately after the personal-site fix.
            full = _scoped_path(folder, file_path)
            if folder:
                scope_item = None
                try:
                    scope_item = _graph(f"/sites/{site_id}/drive/root:/{_enc(folder)}", token)
                except Exception:  # noqa: BLE001 — a folder scope 404s here; that is the normal case
                    scope_item = None
                if scope_item and "file" in scope_item:
                    import posixpath
                    wanted = (file_path or "").strip("/").lower()
                    have = str(scope_item.get("name") or "").lower()
                    # Accept the empty ask, the exact name, or the full path — anything else is
                    # a request for a different file, which this agent's scope does not include.
                    if wanted and wanted != have and posixpath.basename(wanted) != have:
                        return {
                            "error": f"this agent is connected to a single file, \"{scope_item.get('name')}\" — "
                                     f"\"{file_path}\" is outside its scope",
                        }
                    full = folder
            meta = _graph(f"/sites/{site_id}/drive/root:/{_enc(full)}", token)
            size = meta.get("size") or 0
            if size > MAX_BYTES:
                return {"error": f"file is {size} bytes, too large to read inline"}
            blob = _graph(f"/sites/{site_id}/drive/root:/{_enc(full)}:/content", token, raw=True)
        except Exception as e:  # noqa: BLE001
            return {"error": f"SharePoint read failed: {e}"}

        name = (meta.get("name") or file_path).lower()
        try:
            if name.endswith((".txt", ".md", ".csv", ".json", ".log", ".xml", ".yaml", ".yml")):
                text = blob.decode("utf-8", errors="replace")
            elif name.endswith(".pdf"):
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(blob))
                text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
            elif name.endswith(".docx"):
                import docx
                doc = docx.Document(io.BytesIO(blob))
                # .paragraphs alone misses tables entirely — python-docx does not
                # walk them as paragraphs. Confirmed live 2026-08-11: a file whose
                # real content lived in a table returned only its plain-text
                # headings, silently dropping everything else.
                parts = [p.text for p in doc.paragraphs if p.text]
                for i, table in enumerate(doc.tables):
                    parts.append(f"# table {i + 1}")
                    for row in table.rows:
                        # A multi-paragraph cell's OWN text contains \n — join with the
                        # same character as rows and a 10-paragraph cell becomes 9 fake
                        # extra rows, scrambling the table (confirmed live 2026-08-11
                        # against a real file with multi-paragraph "Summary" cells).
                        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                        parts.append(" | ".join(cells))
                text = "\n".join(parts)
            elif name.endswith(".xlsx"):
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
                rows = []
                for ws in wb.worksheets:
                    rows.append(f"# sheet: {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        rows.append(", ".join("" if c is None else str(c) for c in row))
                text = "\n".join(rows)
            else:
                return {"error": f"cannot extract text from '{meta.get('name')}' "
                                 f"— unsupported format. Only text, PDF, Word and Excel are readable."}
        except Exception as e:  # noqa: BLE001
            return {"error": f"text extraction failed for {meta.get('name')}: {e}"}

        truncated = len(text) > MAX_CHARS
        return {
            "file": meta.get("name"),
            "webUrl": meta.get("webUrl"),
            "truncated": truncated,
            "text": text[:MAX_CHARS],
        }

    def sharepoint_list_lists() -> dict:
        """List the SharePoint LISTS and DOCUMENT LIBRARIES on the connected site, so you
        can tell the user what data containers exist before reading from one.

        This is about the site's structure — the lists and libraries themselves — not the
        files inside a folder. Use `sharepoint_list_files` for files.

        Returns:
            dict with `lists` (each: name, displayName, template, webUrl, id) or `error`.
        """
        try:
            token = mint_token(fill)
            site_id, _folder = _resolve_scope(token)
            if not site_id:
                return {"error": "no SharePoint scope configured for this agent"}
            data = _graph(f"/sites/{site_id}/lists", token)
        except Exception as e:  # noqa: BLE001
            return {"error": f"SharePoint list-of-lists failed: {e}"}
        lists = [{
            "name": l.get("name"),
            "displayName": l.get("displayName"),
            "template": (l.get("list") or {}).get("template"),
            "webUrl": l.get("webUrl"),
            "id": l.get("id"),
        } for l in data.get("value", [])]
        return {"site": scope_uri, "lists": lists, "count": len(lists)}

    return [sharepoint_list_files, sharepoint_read_file, sharepoint_list_lists]
