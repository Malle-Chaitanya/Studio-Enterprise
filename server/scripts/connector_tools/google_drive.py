"""Google Drive live tools — named, documented functions instead of the
generic REST fallback, because a bare "call any REST API" tool needs the
model to already know Drive's API conventions by heart (confirmed live
2026-08-10 that it declines to even try). See connectors/confluence.py's
module docstring for the shared build_tools contract.
"""

# create_file/update_file only ever write a RAW TEXT string as the file's bytes —
# there is no code here that generates a real .docx/.xlsx/.pdf binary structure.
# Letting mime_type be set to one of those anyway creates a file that LIES about
# its own format: Drive shows it as a Word document, but the bytes are just text,
# so google_drive_read_file correctly refuses it later as "not a valid docx" —
# permanently broken, self-contradictory, and confusing days after the fact.
# Confirmed live 2026-08-13: a file created this way ("CXXXXXXXXXXXXXXXXXXX.docx")
# had `mimeType: application/vnd...wordprocessingml.document` but its actual 21
# bytes were literally the plain string "hey bob! how are you?". Restrict to types
# a plain string genuinely IS, rather than let the mismatch happen at all.
TEXT_SAFE_MIME_TYPES = {
    "text/plain", "text/markdown", "text/csv", "application/json",
    "text/xml", "application/xml",
}


def build_tools(conn, secret, mint_token, auth_header, fill):
    # Maps a native Google type to (export mimeType, fake filename to dispatch
    # extraction by). Spreadsheets export as .xlsx, not CSV — CSV can only hold
    # ONE sheet, so a multi-tab spreadsheet silently lost every tab but the
    # first (confirmed live 2026-08-11: two tabs with real content came back
    # empty because the first, blank tab was the only one CSV could export).
    GOOGLE_EXPORT_MIME = {
        "application/vnd.google-apps.document": ("text/plain", "export.txt"),
        "application/vnd.google-apps.spreadsheet": (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "export.xlsx",
        ),
        "application/vnd.google-apps.presentation": ("text/plain", "export.txt"),
    }

    def google_drive_list_files(folder_id: str) -> dict:
        """List files and folders inside a Google Drive folder.

        Args:
            folder_id: the Drive folder ID (the part after /folders/ in its
                URL). Use "root" for the top level of My Drive.

        Returns:
            dict with `items` (id, name, mimeType, size, modifiedTime,
            isFolder), `count`, and `truncated` (true if the folder has more
            than MAX_ITEMS files and the list was cut off), or `error`.
        """
        import json as _json
        import urllib.parse
        import urllib.request

        MAX_ITEMS = 1000  # safety cap so one call can't run away on a huge folder
        MAX_PAGES = 20

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        # Auto-paginate internally — handing a raw nextPageToken back to the
        # model/end-user and asking them to paste it in again is bad UX and,
        # confirmed live 2026-08-10, is exactly what happened before this fix:
        # the agent relayed an opaque token string to a human and asked them
        # to repeat their question with it attached.
        items = []
        page_token = ""
        truncated = False
        for _ in range(MAX_PAGES):
            params = {
                "q": f"'{folder_id}' in parents and trashed = false",
                "fields": "nextPageToken, files(id, name, mimeType, size, modifiedTime)",
                "pageSize": "200",
                "supportsAllDrives": "true",
                "includeItemsFromAllDrives": "true",
            }
            if page_token:
                params["pageToken"] = page_token
            url = "https://www.googleapis.com/drive/v3/files?" + urllib.parse.urlencode(params)
            req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
            try:
                with urllib.request.urlopen(req, timeout=25) as resp:
                    data = _json.loads(resp.read().decode("utf-8"))
            except Exception as e:  # noqa: BLE001
                if items:
                    break  # return what we have rather than lose an already-successful partial list
                return {"error": f"Google Drive list failed: {e}"}

            items.extend({
                "id": f.get("id"),
                "name": f.get("name"),
                "mimeType": f.get("mimeType"),
                "size": f.get("size"),
                "modifiedTime": f.get("modifiedTime"),
                "isFolder": f.get("mimeType") == "application/vnd.google-apps.folder",
            } for f in data.get("files", []))

            page_token = data.get("nextPageToken") or ""
            if len(items) >= MAX_ITEMS:
                truncated = bool(page_token)
                items = items[:MAX_ITEMS]
                break
            if not page_token:
                break
        else:
            truncated = bool(page_token)  # MAX_PAGES exhausted with more still remaining

        result = {"folderId": folder_id, "items": items, "count": len(items), "truncated": truncated}
        if truncated:
            result["note"] = f"This folder has more than {len(items)} files — list was capped, not exhaustive."
        return result

    def google_drive_read_file(file_id: str) -> dict:
        """Read the TEXT CONTENT of a Google Drive file, so you can answer
        questions about what it contains.

        Supports plain text, PDF, Word (.docx), Excel (.xlsx), and native
        Google Docs/Sheets/Slides (auto-exported to text/CSV). Images and
        other binary formats cannot be read.

        Args:
            file_id: the Drive file ID (from google_drive_list_files, or the
                /d/<id>/ segment of a shared link).

        Returns:
            dict with `text` (extracted, possibly truncated) or `error`.
        """
        import io
        import json as _json
        import urllib.parse
        import urllib.request

        MAX_BYTES = 20 * 1024 * 1024
        MAX_CHARS = 60000

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        meta_url = (
            "https://www.googleapis.com/drive/v3/files/" + urllib.parse.quote(file_id)
            + "?fields=id,name,mimeType,size,webViewLink&supportsAllDrives=true"
        )
        req = urllib.request.Request(meta_url, headers={"Authorization": f"Bearer {token}"})
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                meta = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive metadata lookup failed: {e}"}

        mime = meta.get("mimeType") or ""
        size = int(meta.get("size") or 0)
        if size > MAX_BYTES:
            return {"error": f"file is {size} bytes, too large to read inline"}

        try:
            if mime in GOOGLE_EXPORT_MIME:
                # Native Google Docs/Sheets/Slides have no raw bytes — Drive
                # rejects alt=media for these ("Use Export with Docs Editors
                # files"). Export converts to a real format first — dispatch
                # extraction below by the EXPORTED format, not the original
                # Google type, so e.g. a spreadsheet's .xlsx export goes
                # through the same multi-sheet-aware openpyxl path a regular
                # uploaded .xlsx file does.
                export_mime, dispatch_name = GOOGLE_EXPORT_MIME[mime]
                dl_url = (
                    "https://www.googleapis.com/drive/v3/files/" + urllib.parse.quote(file_id)
                    + "/export?mimeType=" + urllib.parse.quote(export_mime)
                )
                dl_req = urllib.request.Request(dl_url, headers={"Authorization": f"Bearer {token}"})
                with urllib.request.urlopen(dl_req, timeout=30) as resp:
                    blob = resp.read()
            else:
                dl_url = (
                    "https://www.googleapis.com/drive/v3/files/" + urllib.parse.quote(file_id)
                    + "?alt=media&supportsAllDrives=true"
                )
                dl_req = urllib.request.Request(dl_url, headers={"Authorization": f"Bearer {token}"})
                with urllib.request.urlopen(dl_req, timeout=30) as resp:
                    blob = resp.read()
                dispatch_name = meta.get("name") or ""

            # Dispatch by extension FIRST, but fall back to the real MIME type Drive
            # reported — a file with no extension (e.g. created via
            # google_drive_create_file with a bare name like "A") always fell through
            # to "unsupported format" here even when its actual content was plain
            # text, because every check below only ever looked at the name. Confirmed
            # live 2026-08-13: a file this same tool had just created as text/plain
            # became unreadable by itself one call later, purely because it had no
            # ".txt" suffix.
            name = dispatch_name.lower()
            is_text_mime = mime in ("text/plain", "text/markdown", "text/csv", "application/json", "text/xml", "application/xml")
            is_pdf_mime = mime == "application/pdf"
            is_docx_mime = mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            is_xlsx_mime = mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            if name.endswith((".txt", ".md", ".csv", ".json", ".log", ".xml", ".yaml", ".yml")) or is_text_mime:
                text = blob.decode("utf-8", errors="replace")
            elif name.endswith(".pdf") or is_pdf_mime:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(blob))
                text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
            elif name.endswith(".docx") or is_docx_mime:
                import docx
                doc = docx.Document(io.BytesIO(blob))
                # .paragraphs alone misses tables entirely — python-docx does not
                # walk them as paragraphs. Confirmed live 2026-08-11: a file whose
                # real content lived in a table returned only its plain-text
                # headings, silently dropping everything else.
                parts = [p.text for p in doc.paragraphs if p.text]
                for i, table in enumerate(doc.tables):
                    parts.append(f"# table {i + 1}")
                    for row in table.rows:
                        # A multi-paragraph cell's OWN text contains \n — join with the
                        # same character as rows and a 10-paragraph cell becomes 9 fake
                        # extra rows, scrambling the table (confirmed live 2026-08-11
                        # against a real file with multi-paragraph "Summary" cells).
                        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                        parts.append(" | ".join(cells))
                text = "\n".join(parts)
            elif name.endswith(".xlsx") or is_xlsx_mime:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
                rows = []
                for ws in wb.worksheets:
                    rows.append(f"# sheet: {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        rows.append(", ".join("" if c is None else str(c) for c in row))
                text = "\n".join(rows)
            else:
                return {"error": f"cannot extract text from '{meta.get('name')}' "
                                 f"— unsupported format. Only text, PDF, Word, Excel and "
                                 f"native Google Docs/Sheets/Slides are readable."}
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive read failed for {meta.get('name')}: {e}"}

        truncated = len(text) > MAX_CHARS
        return {
            "file": meta.get("name"),
            "webUrl": meta.get("webViewLink"),
            "truncated": truncated,
            "text": text[:MAX_CHARS],
        }

    def google_drive_get_metadata(file_id: str) -> dict:
        """Get metadata for a Google Drive file or folder — name, type,
        size, modified time, owner and link — WITHOUT reading its content.

        Args:
            file_id: the Drive file or folder ID.

        Returns:
            dict with `id`, `name`, `mimeType`, `size`, `modifiedTime`,
            `webUrl`, `isFolder`, `owners`, or `error`.
        """
        import json as _json
        import urllib.parse
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        url = (
            "https://www.googleapis.com/drive/v3/files/" + urllib.parse.quote(file_id)
            + "?fields=id,name,mimeType,size,modifiedTime,webViewLink,owners&supportsAllDrives=true"
        )
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                meta = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive metadata lookup failed: {e}"}
        return {
            "id": meta.get("id"),
            "name": meta.get("name"),
            "mimeType": meta.get("mimeType"),
            "size": meta.get("size"),
            "modifiedTime": meta.get("modifiedTime"),
            "webUrl": meta.get("webViewLink"),
            "isFolder": meta.get("mimeType") == "application/vnd.google-apps.folder",
            "owners": [o.get("emailAddress") for o in (meta.get("owners") or [])],
        }

    def google_drive_find_by_path(path: str) -> dict:
        """Resolve a human-readable Drive path (e.g. "Reports/2026/budget.xlsx")
        to a file or folder ID. Google Drive has no native path lookup — this
        walks each segment by name starting from the root. Use the returned
        `id` with google_drive_get_metadata or google_drive_read_file to
        actually get metadata or content BY PATH.

        Args:
            path: slash-separated path from the Drive root. Leading/trailing
                slashes are ignored.

        Returns:
            dict with `id`, `name`, `isFolder`, or `error` ("not found" if a
            segment doesn't match, "ambiguous" if more than one item shares
            that name at that level).
        """
        import json as _json
        import urllib.parse
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        segments = [s for s in path.strip("/").split("/") if s]
        if not segments:
            return {"id": "root", "name": "root", "isFolder": True}

        parent_id = "root"
        current = None
        for i, name in enumerate(segments):
            is_last = i == len(segments) - 1
            safe_name = name.replace("\\", "\\\\").replace("'", "\\'")
            q = f"'{parent_id}' in parents and name = '{safe_name}' and trashed = false"
            params = {
                "q": q,
                "fields": "files(id, name, mimeType)",
                "supportsAllDrives": "true",
                "includeItemsFromAllDrives": "true",
            }
            url = "https://www.googleapis.com/drive/v3/files?" + urllib.parse.urlencode(params)
            req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
            try:
                with urllib.request.urlopen(req, timeout=25) as resp:
                    data = _json.loads(resp.read().decode("utf-8"))
            except Exception as e:  # noqa: BLE001
                return {"error": f"Google Drive path lookup failed at '{name}': {e}"}

            matches = data.get("files", [])
            if not matches:
                seen_so_far = "/".join(segments[:i]) or "root"
                return {"error": f"path segment '{name}' not found under '{seen_so_far}'"}
            if len(matches) > 1:
                # Applies to the LAST segment too, not just intermediate folders —
                # confirmed live 2026-08-11: silently picking matches[0] for a
                # duplicate-named FINAL file/folder is exactly as dangerous as doing
                # it for a folder in the middle of the path (e.g. two files both
                # named "X.docx" in the same folder — acting on the wrong one).
                where = "/".join(segments[:i]) or "root"
                return {"error": f"'{name}' is ambiguous under '{where}' — {len(matches)} items share that name; "
                                  f"use google_drive_list_files on that folder to see their distinct IDs"}
            current = matches[0]
            parent_id = current.get("id")

        return {
            "id": current.get("id"),
            "name": current.get("name"),
            "isFolder": current.get("mimeType") == "application/vnd.google-apps.folder",
        }

    def google_drive_search_by_name(name: str, only_folders: bool = False) -> dict:
        """Search ALL of Google Drive for files/folders with an EXACT name —
        use this when you don't already know the full path to something
        (e.g. the user says "the CCB folder" or "X.docx" without saying which
        parent folder it's in). google_drive_find_by_path requires the full
        path already; this does not — it searches everywhere.

        If this returns more than one match, do NOT guess which one is
        right — call google_drive_get_metadata or google_drive_list_files on
        the candidates' parent folders to disambiguate, or ask the user.

        Args:
            name: the exact file or folder name to search for (not a partial
                match).
            only_folders: set true to only match folders — useful for
                resolving an intermediate folder name before searching inside it.

        Returns:
            dict with `matches` (list of {id, name, mimeType, isFolder,
            parents} — `parents` is a list of parent folder IDs, call
            google_drive_get_metadata on one to see its name), `count`, or
            `error`.
        """
        import json as _json
        import urllib.parse
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        safe_name = name.replace("\\", "\\\\").replace("'", "\\'")
        q = f"name = '{safe_name}' and trashed = false"
        if only_folders:
            q += " and mimeType = 'application/vnd.google-apps.folder'"
        params = {
            "q": q,
            "fields": "files(id, name, mimeType, parents)",
            "pageSize": "50",
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
        }
        url = "https://www.googleapis.com/drive/v3/files?" + urllib.parse.urlencode(params)
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                data = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive search failed: {e}"}

        matches = [{
            "id": f.get("id"),
            "name": f.get("name"),
            "mimeType": f.get("mimeType"),
            "isFolder": f.get("mimeType") == "application/vnd.google-apps.folder",
            "parents": f.get("parents") or [],
        } for f in data.get("files", [])]
        return {"matches": matches, "count": len(matches)}

    def google_drive_copy_file(file_id: str, new_name: str = "") -> dict:
        """Copy a Google Drive file. The copy lands in the same folder as
        the original.

        Args:
            file_id: the Drive file ID to copy.
            new_name: optional name for the copy. Defaults to the source
                file's own name.

        Returns:
            dict with the new file's `id`, `name`, `webUrl`, or `error`.
        """
        import json as _json
        import urllib.parse
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        body = {"name": new_name} if new_name else {}
        url = (
            "https://www.googleapis.com/drive/v3/files/" + urllib.parse.quote(file_id)
            + "/copy?fields=id,name,webViewLink&supportsAllDrives=true"
        )
        req = urllib.request.Request(
            url, data=_json.dumps(body).encode("utf-8"),
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                result = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive copy failed: {e}"}
        return {"id": result.get("id"), "name": result.get("name"), "webUrl": result.get("webViewLink")}

    def google_drive_create_file(name: str, folder_id: str, content: str, mime_type: str = "text/plain") -> dict:
        """Create a new file with text content inside a Google Drive folder.

        Args:
            name: the new file's name, e.g. "notes.txt".
            folder_id: the Drive folder ID to create it in. Use "root" for
                the top level of My Drive.
            content: the text content to write into the file.
            mime_type: the file's MIME type. Defaults to plain text.

        Returns:
            dict with the new file's `id`, `name`, `webUrl`, or `error`.
        """
        import json as _json
        import urllib.request

        if mime_type not in TEXT_SAFE_MIME_TYPES:
            return {"error": f"cannot create a real '{mime_type}' file — this tool only writes plain "
                              f"text content, it does not generate a genuine Word/Excel/PDF binary "
                              f"structure. Use a text type instead (e.g. text/plain), or leave "
                              f"mime_type at its default."}

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        # Two steps (create metadata, then upload content) rather than a
        # hand-built multipart/related body — simpler and no extra library.
        # mimeType set HERE too, not just as the upload's Content-Type header below —
        # a name with no recognizable extension (e.g. "A") gives Drive nothing to
        # infer a type from, so the file's stored mimeType could end up generic
        # regardless of what the later upload declares.
        meta_body = _json.dumps({"name": name, "parents": [folder_id], "mimeType": mime_type}).encode("utf-8")
        create_req = urllib.request.Request(
            "https://www.googleapis.com/drive/v3/files?fields=id,name,webViewLink&supportsAllDrives=true",
            data=meta_body,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(create_req, timeout=25) as resp:
                created = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive create failed: {e}"}

        file_id = created.get("id")
        upload_req = urllib.request.Request(
            f"https://www.googleapis.com/upload/drive/v3/files/{file_id}?uploadType=media&supportsAllDrives=true",
            data=content.encode("utf-8"),
            headers={"Authorization": f"Bearer {token}", "Content-Type": mime_type},
            method="PATCH",
        )
        try:
            with urllib.request.urlopen(upload_req, timeout=30):
                pass
        except Exception as e:  # noqa: BLE001
            return {"error": f"file created (id {file_id}) but content upload failed: {e}"}
        return {"id": file_id, "name": created.get("name"), "webUrl": created.get("webViewLink")}

    def _drive_find_folder(token: str, parent_id: str, name: str):
        """Look up ONE folder by exact name under parent_id. Returns its id, or
        None if not found. Shared by google_drive_create_folder and
        google_drive_create_file_by_path so "does this folder already exist"
        is answered the same way in both places."""
        import json as _json
        import urllib.parse
        import urllib.request

        safe_name = name.replace("\\", "\\\\").replace("'", "\\'")
        q = (
            f"'{parent_id}' in parents and name = '{safe_name}' "
            "and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        )
        params = {
            "q": q, "fields": "files(id, name)",
            "supportsAllDrives": "true", "includeItemsFromAllDrives": "true",
        }
        url = "https://www.googleapis.com/drive/v3/files?" + urllib.parse.urlencode(params)
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        with urllib.request.urlopen(req, timeout=25) as resp:
            data = _json.loads(resp.read().decode("utf-8"))
        matches = data.get("files", [])
        return matches[0]["id"] if matches else None

    def google_drive_create_folder(name: str, parent_id: str = "root") -> dict:
        """Create a new, empty folder in Google Drive. Does nothing to any
        files — use google_drive_create_file or
        google_drive_create_file_by_path to also write a file.

        Args:
            name: the new folder's name.
            parent_id: the Drive folder ID to create it in. Defaults to
                "root" (top level of My Drive).

        Returns:
            dict with the new folder's `id`, `name`, `webUrl`, or `error`.
        """
        import json as _json
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        body = _json.dumps({
            "name": name, "parents": [parent_id],
            "mimeType": "application/vnd.google-apps.folder",
        }).encode("utf-8")
        req = urllib.request.Request(
            "https://www.googleapis.com/drive/v3/files?fields=id,name,webViewLink&supportsAllDrives=true",
            data=body,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                result = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive folder creation failed: {e}"}
        return {"id": result.get("id"), "name": result.get("name"), "webUrl": result.get("webViewLink")}

    def google_drive_create_file_by_path(path: str, content: str, mime_type: str = "text/plain") -> dict:
        """Create a file with text content at a Drive PATH, creating any
        missing folders along the way (like "mkdir -p" then write the
        file) — all in this one call. Prefer this over
        google_drive_create_file whenever the destination is given as a
        path (e.g. "Reports/2026/Drafts/notes.txt") rather than a known
        folder ID, especially if a subfolder in that path might not exist
        yet.

        Args:
            path: full path INCLUDING the file name, e.g.
                "Reports/2026/notes.txt". Use just the file name (e.g.
                "notes.txt") to create directly in root.
            content: the text content to write into the file.
            mime_type: the file's MIME type. Defaults to plain text.

        Returns:
            dict with the new file's `id`, `name`, `webUrl`,
            `foldersCreated` (names of any folders that had to be created
            along the way — empty if they all already existed), or `error`.
        """
        import json as _json
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        segments = [s for s in path.strip("/").split("/") if s]
        if not segments:
            return {"error": "path must include at least a file name"}
        file_name = segments[-1]
        folder_names = segments[:-1]

        parent_id = "root"
        folders_created = []
        try:
            for name in folder_names:
                found_id = _drive_find_folder(token, parent_id, name)
                if found_id:
                    parent_id = found_id
                    continue
                made = google_drive_create_folder(name, parent_id)
                if made.get("error"):
                    return {"error": f"could not create folder '{name}': {made['error']}"}
                parent_id = made["id"]
                folders_created.append(name)
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive folder resolution failed: {e}"}

        created = google_drive_create_file(file_name, parent_id, content, mime_type)
        if created.get("error"):
            return created
        created["foldersCreated"] = folders_created
        return created

    def google_drive_update_file(file_id: str, content: str = "", new_name: str = "", mime_type: str = "text/plain") -> dict:
        """Update an existing Google Drive file: replace its content,
        rename it, or both in one call.

        Args:
            file_id: the Drive file ID to update.
            content: new text content to REPLACE the file's current
                content with (does not append). Leave empty to leave
                content unchanged.
            new_name: new name for the file. Leave empty to leave the
                name unchanged.
            mime_type: the content's MIME type, used only if content is
                given. Defaults to plain text.

        Returns:
            dict with `id`, `name`, `webUrl`, or `error`.
        """
        import json as _json
        import urllib.request

        if not content and not new_name:
            return {"error": "nothing to update — provide content and/or new_name"}
        if content and mime_type not in TEXT_SAFE_MIME_TYPES:
            return {"error": f"cannot write '{mime_type}' content — this tool only writes plain text, "
                              f"it does not generate a genuine Word/Excel/PDF binary structure. Use a "
                              f"text type instead (e.g. text/plain), or leave mime_type at its default."}

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        result = {}
        if new_name:
            # Renaming is metadata (JSON PATCH), completely separate from the
            # media-upload PATCH content uses — two different request shapes,
            # so a rename+content update is two real API calls, not one.
            meta_body = _json.dumps({"name": new_name}).encode("utf-8")
            meta_req = urllib.request.Request(
                f"https://www.googleapis.com/drive/v3/files/{file_id}"
                "?fields=id,name,webViewLink&supportsAllDrives=true",
                data=meta_body,
                headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                method="PATCH",
            )
            try:
                with urllib.request.urlopen(meta_req, timeout=25) as resp:
                    result = _json.loads(resp.read().decode("utf-8"))
            except Exception as e:  # noqa: BLE001
                return {"error": f"Google Drive rename failed: {e}"}

        if content:
            url = (
                f"https://www.googleapis.com/upload/drive/v3/files/{file_id}"
                "?uploadType=media&fields=id,name,webViewLink&supportsAllDrives=true"
            )
            req = urllib.request.Request(
                url, data=content.encode("utf-8"),
                headers={"Authorization": f"Bearer {token}", "Content-Type": mime_type},
                method="PATCH",
            )
            try:
                with urllib.request.urlopen(req, timeout=30) as resp:
                    result = _json.loads(resp.read().decode("utf-8"))
            except Exception as e:  # noqa: BLE001
                partial = {"error": f"Google Drive content update failed: {e}"}
                if new_name:
                    partial["note"] = f"rename to '{new_name}' already succeeded before this failed"
                return partial

        return {"id": result.get("id"), "name": result.get("name"), "webUrl": result.get("webViewLink")}

    def google_drive_delete_file(file_id: str) -> dict:
        """Move a Google Drive file or folder to Trash. This is RECOVERABLE
        (from Drive's Trash) — it does not permanently erase the file.

        Args:
            file_id: the Drive file or folder ID to trash.

        Returns:
            dict with `trashed: true` and the `id`, or `error`.
        """
        import json as _json
        import urllib.request

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        url = (
            "https://www.googleapis.com/drive/v3/files/" + file_id
            + "?fields=id,trashed&supportsAllDrives=true"
        )
        body = _json.dumps({"trashed": True}).encode("utf-8")
        req = urllib.request.Request(
            url, data=body,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            method="PATCH",
        )
        try:
            with urllib.request.urlopen(req, timeout=25) as resp:
                result = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive trash failed: {e}"}
        return {"trashed": bool(result.get("trashed")), "id": result.get("id")}

    def google_drive_extract_archive(file_id: str, dest_folder_id: str) -> dict:
        """Extract a .zip file from Google Drive into a destination folder,
        uploading each entry as its own new file. Only .zip is supported —
        Drive has no native extract capability, so this downloads the
        archive, unzips it in memory, and re-uploads each entry.

        Args:
            file_id: the Drive file ID of the .zip to extract.
            dest_folder_id: the Drive folder ID to extract into.

        Returns:
            dict with `extracted` (list of {name, id}), `skipped` (list of
            "name: reason" for entries that failed), `count`, or `error`.
        """
        import io
        import json as _json
        import urllib.request
        import zipfile

        MAX_ARCHIVE_BYTES = 50 * 1024 * 1024

        try:
            token = mint_token(fill)
        except Exception as e:  # noqa: BLE001
            return {"error": f"auth failed: {e}"}

        meta_url = (
            "https://www.googleapis.com/drive/v3/files/" + file_id
            + "?fields=name,mimeType,size&supportsAllDrives=true"
        )
        meta_req = urllib.request.Request(meta_url, headers={"Authorization": f"Bearer {token}"})
        try:
            with urllib.request.urlopen(meta_req, timeout=25) as resp:
                meta = _json.loads(resp.read().decode("utf-8"))
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive metadata lookup failed: {e}"}

        if not (meta.get("name") or "").lower().endswith(".zip"):
            return {"error": f"'{meta.get('name')}' is not a .zip file — only zip archives are supported"}
        size = int(meta.get("size") or 0)
        if size > MAX_ARCHIVE_BYTES:
            return {"error": f"archive is {size} bytes, too large to extract inline"}

        dl_url = (
            "https://www.googleapis.com/drive/v3/files/" + file_id + "?alt=media&supportsAllDrives=true"
        )
        dl_req = urllib.request.Request(dl_url, headers={"Authorization": f"Bearer {token}"})
        try:
            with urllib.request.urlopen(dl_req, timeout=30) as resp:
                blob = resp.read()
        except Exception as e:  # noqa: BLE001
            return {"error": f"Google Drive download failed: {e}"}

        try:
            zf = zipfile.ZipFile(io.BytesIO(blob))
        except Exception as e:  # noqa: BLE001
            return {"error": f"could not open archive: {e}"}

        extracted, skipped = [], []
        for info in zf.infolist():
            if info.is_dir():
                continue
            member_name = info.filename.split("/")[-1]
            try:
                member_bytes = zf.read(info)
                meta_body = _json.dumps({"name": member_name, "parents": [dest_folder_id]}).encode("utf-8")
                create_req = urllib.request.Request(
                    "https://www.googleapis.com/drive/v3/files?fields=id,name&supportsAllDrives=true",
                    data=meta_body,
                    headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                    method="POST",
                )
                with urllib.request.urlopen(create_req, timeout=25) as resp:
                    created = _json.loads(resp.read().decode("utf-8"))
                upload_req = urllib.request.Request(
                    "https://www.googleapis.com/upload/drive/v3/files/" + created["id"]
                    + "?uploadType=media&supportsAllDrives=true",
                    data=member_bytes,
                    headers={"Authorization": f"Bearer {token}", "Content-Type": "application/octet-stream"},
                    method="PATCH",
                )
                with urllib.request.urlopen(upload_req, timeout=30):
                    pass
                extracted.append({"name": member_name, "id": created["id"]})
            except Exception as e:  # noqa: BLE001
                skipped.append(f"{member_name}: {e}")

        return {"extracted": extracted, "skipped": skipped, "count": len(extracted)}

    return [
        google_drive_list_files,
        google_drive_read_file,
        google_drive_get_metadata,
        google_drive_find_by_path,
        google_drive_search_by_name,
        google_drive_copy_file,
        google_drive_create_file,
        google_drive_create_folder,
        google_drive_create_file_by_path,
        google_drive_update_file,
        google_drive_delete_file,
        google_drive_extract_archive,
    ]
